from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models.statement_schema import BankStatement

class DOCXGenerator:
    def generate(self, statement: BankStatement, output_path: str):
        """Generate formatted DOCX from BankStatement"""
        doc = Document()
        
        # Title
        title = doc.add_heading(statement.header.bank_name or "Bank Statement", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Account Details
        doc.add_heading('Account Information', level=1)
        
        details_table = doc.add_table(rows=5, cols=2)
        details_table.style = 'Light Grid Accent 1'
        
        details = [
            ("Account Holder:", statement.header.account_holder),
            ("Account Number:", statement.header.account_number),
            ("IFSC Code:", statement.header.ifsc or "N/A"),
            ("MICR Code:", statement.header.micr or "N/A"),
            ("Branch:", statement.header.branch or "N/A"),
        ]
        
        for i, (label, value) in enumerate(details):
            row = details_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Transactions
        doc.add_heading('Transaction Details', level=1)
        
        txn_table = doc.add_table(rows=1, cols=5)
        txn_table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = txn_table.rows[0].cells
        headers = ["Date", "Description", "Credit", "Debit", "Balance"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Data rows
        for txn in statement.transactions:
            row_cells = txn_table.add_row().cells
            row_cells[0].text = txn.date.strftime("%d-%m-%Y")
            row_cells[1].text = txn.description
            row_cells[2].text = f"₹{txn.credit:.2f}" if txn.credit > 0 else "-"
            row_cells[3].text = f"₹{txn.debit:.2f}" if txn.debit > 0 else "-"
            row_cells[4].text = f"₹{txn.balance:.2f}"
        
        # Summary row
        summary_row = txn_table.add_row().cells
        summary_row[0].text = "TOTAL"
        summary_row[0].paragraphs[0].runs[0].font.bold = True
        summary_row[2].text = f"₹{statement.total_credits:.2f}"
        summary_row[3].text = f"₹{statement.total_debits:.2f}"
        summary_row[4].text = f"₹{statement.closing_balance:.2f}"
        
        doc.save(output_path)
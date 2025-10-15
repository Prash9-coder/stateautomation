from docx import Document
from parsers.llm_extractor import LLMExtractor
from models.statement_schema import BankStatement

class DOCXParser:
    def __init__(self):
        self.llm_extractor = LLMExtractor()
    
    def parse(self, file_path: str) -> BankStatement:
        """Parse DOCX bank statement"""
        doc = Document(file_path)
        
        # Extract all text
        text = "\n".join([para.text for para in doc.paragraphs])
        
        # Extract tables
        for table in doc.tables:
            table_text = "\n".join([
                "\t".join([cell.text for cell in row.cells])
                for row in table.rows
            ])
            text += "\n" + table_text
        
        # Use LLM to structure the data
        extracted_data = self.llm_extractor.extract_structured_data(text)
        statement = BankStatement(**extracted_data)
        
        return statement